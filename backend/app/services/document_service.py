"""详情页文档导出：Word / PDF / HTML。

使用图片路径或 base64 嵌入；保证离线可用。
"""
import base64
from io import BytesIO
from pathlib import Path
from typing import List

from PIL import Image
from docx import Document
from docx.shared import Inches, Pt
from jinja2 import Template


HTML_TEMPLATE = Template("""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{{ project.name }} - 详情页</title>
<style>
  * { box-sizing: border-box; }
  body { margin: 0; font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", "PingFang SC", "Hiragino Sans GB", sans-serif; background: #f7f7f7; color: #222; }
  .cover { padding: 80px 40px; background: linear-gradient(135deg, #4f6df5, #6c8cff); color: #fff; text-align: center; }
  .cover h1 { font-size: 36px; margin: 0 0 12px; }
  .cover .meta { opacity: 0.85; font-size: 14px; }
  .info { padding: 24px 40px; background: #fff; margin: 16px; border-radius: 8px; box-shadow: 0 1px 4px rgba(0,0,0,.04); }
  .info dl { display: grid; grid-template-columns: 120px 1fr; gap: 8px 16px; margin: 0; }
  .info dt { color: #888; }
  .module { background: #fff; margin: 16px; border-radius: 8px; overflow: hidden; box-shadow: 0 1px 4px rgba(0,0,0,.04); }
  .module h2 { margin: 0; padding: 16px 24px; font-size: 18px; background: #f0f3ff; color: #4f6df5; }
  .module .desc { padding: 0 24px 12px; color: #666; font-size: 13px; }
  .module img { width: 100%; display: block; }
  .footer { text-align: center; padding: 24px; color: #999; font-size: 12px; }
</style>
</head>
<body>
  <div class="cover">
    <h1>{{ project.name }}</h1>
    <div class="meta">
      {{ project.industry }} · {{ project.target_market or '通用市场' }} · {{ project.target_platform or '通用平台' }} · {{ project.language }}
    </div>
  </div>
  <div class="info">
    <dl>
      <dt>产品名称</dt><dd>{{ project.product_name or '—' }}</dd>
      <dt>视觉风格</dt><dd>{{ project.visual_style or '—' }}</dd>
      <dt>分辨率 / 比例</dt><dd>{{ project.resolution }} / {{ project.aspect_ratio }}</dd>
      <dt>核心卖点</dt><dd>{{ project.product_selling_points or '—' }}</dd>
      <dt>目标用户</dt><dd>{{ project.product_target_audience or '—' }}</dd>
    </dl>
  </div>
  {% for m in modules %}
  <div class="module">
    <h2>{{ loop.index }}. {{ m.name_zh }}</h2>
    {% if m.desc_zh %}<div class="desc">{{ m.desc_zh }}</div>{% endif %}
    {% for img in m.images %}
      <img src="{{ img.src }}" alt="{{ m.name_zh }}">
    {% endfor %}
    {% if not m.images %}<div class="desc" style="color:#aaa">（暂无图片）</div>{% endif %}
  </div>
  {% endfor %}
  <div class="footer">由 AI 电商详情页工作台 生成 · {{ now }}</div>
</body>
</html>
""")


def _read_image_b64(path: Path) -> str:
    if not path.exists():
        return ""
    suffix = path.suffix.lower()
    mime = "image/png" if suffix == ".png" else "image/jpeg"
    return f"data:{mime};base64," + base64.b64encode(path.read_bytes()).decode("ascii")


def _group_assets_by_module(modules: List[dict], assets) -> List[dict]:
    """组装 [{key, name_zh, desc_zh, images: [{src, caption}]}]。"""
    grouped = []
    asset_index = {}
    for a in assets:
        asset_index.setdefault(a.module_key, []).append(a)

    name_map = {m["key"]: m["name_zh"] for m in _ALL_MODULE_META}
    desc_map = {m["key"]: m["desc_zh"] for m in _ALL_MODULE_META}

    for m in modules:
        key = m.get("key")
        imgs = []
        for a in asset_index.get(key, []):
            if a.status != "success" or not a.file_path:
                continue
            src = _read_image_b64(Path(a.file_path))
            if src:
                imgs.append({"src": src, "caption": a.prompt or ""})
        grouped.append({
            "key": key,
            "name_zh": m.get("name_zh") or name_map.get(key, key),
            "desc_zh": m.get("desc_zh") or desc_map.get(key, ""),
            "images": imgs,
        })
    return grouped


# 复用 prompts 的模块元数据
from app.utils.prompts import COMMON_MODULES as _ALL_MODULE_META  # noqa: E402


def export_html(project, assets) -> Path:
    grouped = _group_assets_by_module(project.module_plan or [], assets)
    html = HTML_TEMPLATE.render(
        project=project,
        modules=grouped,
        now=datetime.now().strftime("%Y-%m-%d %H:%M"),
    )
    out = Path(project.workdir) / "09_导出" / f"{project.name}_详情页.html"
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text(html, encoding="utf-8")
    return out


def export_docx(project, assets) -> Path:
    doc = Document()
    doc.add_heading(project.name, level=0)
    doc.add_paragraph(
        f"行业: {project.industry}    市场: {project.target_market}    平台: {project.target_platform}    语言: {project.language}"
    )
    doc.add_paragraph(
        f"产品: {project.product_name or '—'}    分辨率: {project.resolution}    比例: {project.aspect_ratio}"
    )
    if project.product_selling_points:
        doc.add_paragraph(f"核心卖点: {project.product_selling_points}")
    if project.product_target_audience:
        doc.add_paragraph(f"目标用户: {project.product_target_audience}")

    asset_index = {}
    for a in assets:
        asset_index.setdefault(a.module_key, []).append(a)

    name_map = {m["key"]: m["name_zh"] for m in _ALL_MODULE_META}

    for i, m in enumerate(project.module_plan or [], 1):
        key = m.get("key")
        title = m.get("name_zh") or name_map.get(key, key)
        doc.add_heading(f"{i}. {title}", level=1)
        for a in asset_index.get(key, []):
            if a.status != "success" or not a.file_path:
                continue
            p = Path(a.file_path)
            if p.exists():
                try:
                    doc.add_picture(str(p), width=Inches(6))
                except Exception:
                    doc.add_paragraph(f"[图片加载失败] {p}")
            if a.prompt:
                doc.add_paragraph(f"Prompt: {a.prompt}")

    out = Path(project.workdir) / "09_导出" / f"{project.name}_详情页.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


def export_pdf(project, assets) -> Path:
    """简化版：先导 HTML，再用 weasyprint 渲染 PDF。"""
    try:
        from weasyprint import HTML
    except Exception as e:
        # 退而求其次：保存 HTML 并返回 html 路径
        html_path = export_html(project, assets)
        raise RuntimeError(
            f"weasyprint 不可用 ({e})，已导出 HTML 版本，请用浏览器打印 PDF: {html_path}"
        )
    html_path = export_html(project, assets)
    out = Path(project.workdir) / "09_导出" / f"{project.name}_详情页.pdf"
    HTML(string=html_path.read_text(encoding="utf-8")).write_pdf(str(out))
    return out


# 引入 datetime 缺失
from datetime import datetime  # noqa: E402
